from pathlib import Path

from pypdf import PdfReader
from docx import Document as DocxDocument

from backend.core.ingestion.base import BaseParser


class ParserError(Exception):
    pass


class EmptyPDFError(ParserError):
    pass


class PDFParser(BaseParser):
    def parse(self, file_path: str) -> str:
        path = Path(file_path)
        if not path.exists():
            raise ParserError(f"Файл не найден: {file_path}")

        try:
            reader = PdfReader(str(path))
        except Exception as exc:
            raise ParserError(f"Не удалось открыть PDF: {exc}") from exc

        pages_text = [
            text for page in reader.pages if (text := page.extract_text() or "").strip()
        ]
        full_text = "\n\n".join(pages_text).strip()
        if not full_text:
            raise EmptyPDFError(
                f"PDF не содержит извлекаемого текста (возможно, это скан): {file_path}"
            )
        return full_text


class TextParser(BaseParser):
    def parse(self, file_path: str) -> str:
        path = Path(file_path)
        if not path.exists():
            raise ParserError(f"Файл не найден: {file_path}")
        return path.read_text(encoding="utf-8-sig", errors="replace")


class DocxParser(BaseParser):
    def parse(self, file_path: str) -> str:
        path = Path(file_path)
        if not path.exists():
            raise ParserError(f"Файл не найден: {file_path}")

        try:
            doc = DocxDocument(str(path))
        except Exception as exc:
            raise ParserError(f"Не удалось открыть DOCX: {exc}") from exc

        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)
        return "\n\n".join(paragraphs).strip()


PARSER_REGISTRY: dict[str, type[BaseParser]] = {
    ".pdf": PDFParser,
    ".txt": TextParser,
    ".docx": DocxParser,
}

SUPPORTED_EXTENSIONS: tuple[str, ...] = tuple(PARSER_REGISTRY.keys())


def get_parser_for(filename: str) -> BaseParser | None:
    ext = Path(filename).suffix.lower()
    cls = PARSER_REGISTRY.get(ext)
    return cls() if cls else None
